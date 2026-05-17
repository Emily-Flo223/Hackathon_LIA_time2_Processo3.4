"""
Gerador de Relatórios Sintéticos — Processo 3.4 PROECE/UFMS
Hackathon Meta 3 · UFMS Apoia MDA

Gera 30 relatórios de ações de extensão como arquivos reais:
  - 15 em formato DOCX  (modelo estilo SIGProj/UFMS — estruturado em tabelas)
  - 15 em formato PDF   (modelo estilo IFMG — discursivo com seções numeradas)

Uso:
  python gerar_relatorios_extensao.py [--seed 42] [--out ./relatorios]
"""

import argparse
import random
import csv
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from datetime import date, timedelta
from pathlib import Path

# ── DOCX ──────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── PDF ───────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY


# ══════════════════════════════════════════════════════════════════════════
# DADOS SINTÉTICOS
# ══════════════════════════════════════════════════════════════════════════

UNIDADES = ["INQUI", "FACOM", "CCHS", "CCBS", "FAALC", "CPAQ", "FAENG"]

TITULOS = [
    "Grupo Arandú de Tecnologias e Ensino de Ciências (GATEC)",
    "Oficinas de Ciências Multidisciplinares – Officiências",
    "Letramento Digital para Comunidades Rurais do MS",
    "Saúde Bucal nas Escolas Públicas de Campo Grande",
    "Hortas Comunitárias e Segurança Alimentar",
    "Robótica Educacional no Ensino Médio",
    "Empreendedorismo Jovem no Pantanal",
    "Biodiversidade do Cerrado: Educação Ambiental",
    "Inclusão Digital para Idosos – CampuSênior",
    "Apoio Psicossocial às Comunidades Indígenas de MS",
    "Contação de Histórias e Incentivo à Leitura",
    "Monitoramento Hídrico Participativo do Rio Paraguai",
    "Arte e Cultura: Oficinas de Teatro para Jovens",
    "Consultoria Tributária Gratuita para MEIs",
    "Ciência Cidadã: Meteorologia Escolar",
    "Rede Maker Pantaneira: Inovação e Prototipagem Social",
    "Educação Financeira para Famílias em Vulnerabilidade Social",
    "Energia Solar nas Escolas do Mato Grosso do Sul",
    "Mulheres na Tecnologia: Capacitação em Programação e Dados",
    "Agroecologia Urbana e Compostagem Comunitária",
    "Alfabetização Científica para Crianças do Ensino Fundamental",
    "Matemática Aplicada ao Cotidiano Escolar",
    "Laboratório Itinerante de Física Experimental",
    "Capacitação em Primeiros Socorros para Comunidades Rurais",
    "Observatório Jovem de Sustentabilidade e Clima",
    "Música, Inclusão e Desenvolvimento Social",
    "Oficina de Desenvolvimento de Jogos Educacionais",
    "Tecnologias Assistivas para Inclusão Escolar",
    "Memória Cultural e Patrimônio das Comunidades Tradicionais",
    "Capacitação em Inteligência Artificial para Educação Básica"
]

NOMES = [
    "Ana Souza", "Bruno Lima", "Carla Mendes", "Diego Oliveira",
    "Elisa Ferreira", "Fábio Santos", "Gabriela Costa", "Henrique Alves",
    "Isabela Martins", "João Pereira", "Karla Rocha", "Lucas Teixeira",
    "Mariana Campos", "Neto Barbosa", "Olivia Nunes", "Paulo Cardoso",
    "Rafael Gomes", "Sabrina Lopes", "Thiago Vieira", "Zélia Borges",
    "Vítor Dias", "Yara Fernandes", "Wagner Ribeiro", "Ximena Castro", "Valentina Moreira",
    "Zeca Araújo","Adriana Pires", "Breno Figueiredo", "Clara Monteiro", "Danilo Correia",
    "Eduarda Silva", "Felipe Rocha", "Giovana Almeida", "Hugo Costa", "Isadora Martins",
    "Jorge Santos", "Kelly Oliveira", "Leonardo Lima", "Marina Ferreira", "Nicolas Barbosa",
    "Patrícia Gomes", "Quirino Teixeira", "Renata Dias", "Samuel Vieira", "Tatiana Lopes",
    "Ursula Santos", "Victor Almeida", "Wanda Ribeiro", "Xavier Mendes", "Yasmin Costa",
    "Zuleica Nunes", "André Pires", "Bianca Figueiredo", "César Monteiro", "Débora Correia",
    "Eduardo Silva", "Fernanda Rocha", "Gustavo Almeida", "Helena Costa", "Igor Martins",
    "Júlia Santos", "Kleber Oliveira", "Larissa Lima", "Marcos Ferreira", "Nathalia Barbosa",
    "Patrício Gomes", "Quésia Teixeira", "Renan Dias", "Sofia Vieira", "Thiara Lopes",
    "Ulisses Santos", "Vanessa Almeida", "Wellington Ribeiro", "Xuxa Mendes", "Yuri Costa",
    "Zoraide Nunes", "Alex Pires", "Bruna Figueiredo", "Carlos Monteiro", "Diana Correia",
]

AREAS = [
    "Ciências Exatas e da Terra", "Ciências Biológicas", "Engenharia",
    "Ciências da Saúde", "Ciências Humanas", "Ciências Sociais Aplicadas",
]

LINHAS = [
    "Alfabetização, leitura e escrita", "Educação profissional",
    "Meio ambiente", "Saúde humana", "Tecnologia e produção",
    "Desenvolvimento humano", "Arte, cultura e diversidade", "Cidadania e direitos humanos",
    "Inclusão social e digital", "Agricultura familiar e segurança alimentar",
    "Comunicação e divulgação científica", "Esporte e lazer",
    "Gestão pública e políticas sociais", "Inovação social e tecnológica",
]

ODS_LISTA = [
    "ODS 4 – Educação de qualidade",
    "ODS 5 – Igualdade de gênero",
    "ODS 10 – Reducao das desigualdades",
    "ODS 3 – Saúde e bem-estar",
    "ODS 8 – Trabalho decente",
    "ODS 9 – Inovacao e infraestrutura",
    "ODS 11 – Cidades e comunidades sustentáveis",
    "ODS 13 – Ação contra a mudança global do clima",
    "ODS 15 – Vida terrestre",
    "ODS 17 – Parcerias e meios de implementação",
    "ODS 1 – Erradicação da pobreza",
    "ODS 2 – Fome zero e agricultura sustentável",
    "ODS 6 – Água potável e saneamento",
    "ODS 7 – Energia acessível e limpa",
    "ODS 12 – Consumo e produção responsáveis",
    "ODS 16 – Paz, justiça e instituições eficazes",
    "ODS 14 – Vida na água",
    "ODS 5 – Igualdade de gênero",
    "ODS 4 – Educação de qualidade",
    "ODS 17 – Parcerias e meios de implementação",
    "ODS 3 – Saúde e bem-estar",
]

TIPOS_VINCULO = ["PIBEXT", "PIVEXT"]

ERROS = {
    "OK": (10, "Sem erros"),
    "E1": (5,  "Valor financeiro declarado diverge do concedido > 1%"),
    "E2": (4,  "Bolsista declarado nao consta na base oficial"),
    "E3": (3,  "Periodo do bolsista fora do vinculo formal"),
    "E4": (4,  "Secao obrigatoria ausente"),
    "E5": (2,  "Total de horas inconsistente com soma das atividades"),
    "E6": (2,  "Tipo de relatorio inadequado para a fase do projeto"),
}

DECISAO = {
    "OK": "APROVAR",
    "E1": "DEVOLVER PARA AJUSTES",
    "E2": "DEVOLVER PARA AJUSTES",
    "E3": "DEVOLVER PARA AJUSTES",
    "E4": "DEVOLVER PARA REELABORACAO",
    "E5": "DEVOLVER PARA AJUSTES",
    "E6": "DEVOLVER PARA AJUSTES",
}


# ══════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════

def fmt_data(d: date) -> str:
    return d.strftime("%d/%m/%Y")

def data_entre(rng, inicio: date, fim: date) -> date:
    delta = (fim - inicio).days
    return inicio + timedelta(days=rng.randint(0, max(delta, 0)))

def cpf_fake(rng) -> str:
    d = [rng.randint(0, 9) for _ in range(11)]
    s = "".join(map(str, d))
    return f"{s[:3]}.{s[3:6]}.{s[6:9]}-{s[9:]}"

def rga_fake(rng) -> str:
    return f"202{rng.randint(2,6)}{rng.randint(10000000, 99999999)}"

def protocolo(rng, unidade) -> str:
    return f"{unidade}.{rng.randint(100000, 999999)}"

def valor_r(rng, mn, mx) -> float:
    return round(rng.uniform(mn, mx), 2)


# ══════════════════════════════════════════════════════════════════════════
# GERAÇÃO DOS DADOS
# ══════════════════════════════════════════════════════════════════════════

def gerar_dados(rng, idx, tipo_erro):
    unidade = rng.choice(UNIDADES)
    titulo  = rng.choice(TITULOS)
    coord   = rng.choice(NOMES)
    prot    = protocolo(rng, unidade)
    area    = rng.choice(AREAS)
    linha   = rng.choice(LINHAS)
    ods     = rng.sample(ODS_LISTA, rng.randint(1, 3))

    inicio_proj = data_entre(rng, date(2023, 3, 1), date(2024, 3, 1))
    fim_proj    = inicio_proj + timedelta(days=rng.randint(180, 730))

    # E6: relatório final enviado antes do projeto terminar
    if tipo_erro == "E6":
        tipo_rel   = "Final"
        data_envio = data_entre(rng, inicio_proj, inicio_proj + timedelta(days=45))
    else:
        tipo_rel = rng.choice(["Parcial", "Final"])
        if tipo_rel == "Final":
            data_envio = data_entre(rng, fim_proj, fim_proj + timedelta(days=60))
        else:
            data_envio = data_entre(rng, inicio_proj + timedelta(days=60), fim_proj)

    valor_concedido = valor_r(rng, 2000.0, 25000.0)
    tem_recurso     = rng.random() > 0.2

    if not tem_recurso:
        valor_declarado = 0.0
    elif tipo_erro == "E1":
        desvio = rng.choice([-1, 1]) * rng.uniform(0.05, 0.20)
        valor_declarado = round(valor_concedido * (1 + desvio), 2)
    else:
        valor_declarado = round(valor_concedido * rng.uniform(0.995, 1.005), 2)

    # Bolsistas
    n_bol = rng.randint(1, 3)
    bolsistas = []
    for i in range(n_bol):
        ini_bol = data_entre(rng, date(2024, 1, 1), date(2024, 6, 30))
        fim_bol = data_entre(rng, ini_bol + timedelta(days=90), date(2025, 12, 31))
        obs = ""
        cpf_bol = cpf_fake(rng)

        if tipo_erro == "E2" and i == 0:
            cpf_bol = cpf_fake(rng)  # CPF que nao existira na base
            obs = "[CPF NAO LOCALIZADO NA BASE OFICIAL]"
        elif tipo_erro == "E3" and i == 0:
            ini_bol = date(2022, 1, 15)
            fim_bol = date(2022, 8, 30)
            obs = "[PERIODO FORA DO VINCULO FORMAL]"

        bolsistas.append({
            "nome":    rng.choice(NOMES),
            "cpf":     cpf_bol,
            "vinculo": rng.choice(TIPOS_VINCULO),
            "inicio":  fmt_data(ini_bol),
            "fim":     fmt_data(fim_bol),
            "obs":     obs,
        })

    # Atividades
    nomes_atv = [
        "Oficina de ciencias experimentais",
        "Visita tecnica as escolas parceiras",
        "Minicurso de programacao",
        "Palestra de divulgacao cientifica",
        "Roda de conversa com a comunidade",
        "Elaboracao de material didatico",
        "Capacitacao de monitores",
        "Feira de ciencias escolar",
        "Expedicao cientifica de campo",
    ]
    n_atv = rng.randint(3, 6)
    atividades = []
    for nome_atv in rng.sample(nomes_atv, n_atv):
        atividades.append({
            "descricao": nome_atv,
            "data":      fmt_data(data_entre(rng, date(2024, 2, 1), date(2025, 2, 1))),
            "local":     rng.choice(["UFMS – Campus CG", "Escola parceira", "Comunidade local"]),
            "part_ext":  rng.randint(10, 200),
            "horas":     rng.randint(4, 40),
        })
    soma_horas  = sum(a["horas"] for a in atividades)
    total_horas = soma_horas + (rng.choice([-60, 80]) if tipo_erro == "E5" else 0)

    pub_int = {
        "Estudantes de graduacao UFMS":    rng.randint(5, 60),
        "Estudantes pos-graduacao UFMS":   rng.randint(0, 10),
        "Professores UFMS":                rng.randint(1, 5),
        "Tecnicos UFMS":                   rng.randint(0, 3),
    }
    pub_ext = {
        "Estudantes educacao basica":      rng.randint(0, 500),
        "Profissionais da Educacao":       rng.randint(0, 50),
        "Profissionais outras areas":      rng.randint(0, 100),
        "Comunidade em geral":             rng.randint(0, 200),
    }

    secoes = {
        "sintese": (
            f"Durante o periodo de referencia, o projeto {titulo} executou "
            f"{n_atv} atividades, impactando aproximadamente "
            f"{sum(a['part_ext'] for a in atividades)} participantes externos. "
            "As acoes foram desenvolvidas em parceria com escolas publicas e "
            "organizacoes comunitarias, mantendo registro sistematico de todas as etapas."
        ),
        "objetivos_geral": (
            f"Promover a articulacao entre a universidade e a comunidade por meio de "
            f"acoes de extensao vinculadas a area de {area}, contribuindo para a "
            "formacao integral dos estudantes e para o desenvolvimento regional."
        ),
        "objetivos_especificos": [
            "Realizar oficinas e atividades praticas junto ao publico-alvo.",
            "Produzir materiais didaticos adaptados a realidade local.",
            f"Fortalecer a parceria institucional no ambito da linha '{linha}'.",
            "Registrar e divulgar os resultados em eventos cientificos.",
        ],
        "metodologia": (
            "As atividades foram planejadas em ciclos quinzenais com reunioes de "
            "alinhamento entre os membros da equipe. Foram utilizadas metodologias "
            "ativas de aprendizagem, como aprendizagem baseada em projetos e rodas "
            "de conversa. Os dados foram sistematizados em planilhas compartilhadas."
        ),
        "resultados": (
            "Foram alcancados os objetivos previstos, com superacao das metas "
            "quantitativas estabelecidas no edital. O indicador de satisfacao dos "
            "participantes externos atingiu 87% de avaliacoes positivas, medido por "
            "questionario aplicado ao final das atividades."
        ),
        "contribuicoes": (
            "A execucao do projeto evidenciou a necessidade de aprimoramento continuo "
            "das metodologias de ensino adotadas. Elementos identificados poderao "
            "ser explorados por pesquisas futuras e novas acoes de extensao."
        ),
        "dificuldades": rng.choice([
            "Dificuldade de deslocamento as comunidades rurais em periodo de chuvas.",
            "Rotatividade de estudantes bolsistas ao longo do projeto.",
            "Atraso na liberacao dos recursos de custeio, impactando duas atividades.",
            "Nenhuma dificuldade relevante foi identificada no periodo.",
        ]),
        "consideracoes": (
            f"O projeto {titulo} contribuiu de forma efetiva para a formacao integral "
            "dos discentes envolvidos, promovendo a articulacao entre ensino, pesquisa "
            "e extensao. Recomenda-se a continuidade e ampliacao das acoes."
        ),
        "referencias": [
            "BRASIL. Lei no 9.394/1996 – Lei de Diretrizes e Bases da Educacao Nacional.",
            "FREIRE, Paulo. Pedagogia do Oprimido. 17. ed. Rio de Janeiro: Paz e Terra, 1987.",
            "FORPROEX. Politica Nacional de Extensao Universitaria. Manaus, 2012.",
            "UFMS. Resolucao no 291/2019 – Regulamento das Acoes de Extensao.",
        ],
    }

    secao_removida = None
    if tipo_erro == "E4":
        secao_removida = rng.choice(["resultados", "metodologia", "consideracoes", "sintese"])
        secoes[secao_removida] = ""

    produtos = [{"tipo": "Publicacao/Divulgacao em midias", "titulo": titulo, "ano": "2024"}]
    if rng.random() > 0.5:
        produtos.append({"tipo": "Material didatico produzido",
                          "titulo": f"Apostila – {titulo}", "ano": "2024"})

    prestacao_contas = []
    if valor_declarado > 0:
        num_gastos = rng.randint(2, 4)
        restante = valor_declarado
        for i in range(num_gastos - 1):
            gasto = round(rng.uniform(0.1, 0.4) * valor_declarado, 2)
            prestacao_contas.append({
                "item": rng.choice(["Material de consumo", "Serviços de terceiros", "Equipamentos", "Passagens e Despesas com Locomoção"]),
                "valor": gasto
            })
            restante -= gasto
        prestacao_contas.append({
            "item": rng.choice(["Material de consumo", "Bolsas", "Serviços de terceiros"]),
            "valor": round(restante, 2)
        })

    return {
        "idx":              idx,
        "tipo_erro":        tipo_erro,
        "decisao_esperada": DECISAO[tipo_erro],
        "id_relatorio":     str(rng.randint(80000000, 99999999)),
        "protocolo":        prot,
        "titulo":           titulo,
        "unidade":          unidade,
        "area":             area,
        "linha":            linha,
        "ods":              ods,
        "coordenador":      coord,
        "cpf_coord":        cpf_fake(rng),
        "tipo_relatorio":   tipo_rel,
        "inicio_projeto":   fmt_data(inicio_proj),
        "fim_projeto":      fmt_data(fim_proj),
        "data_envio":       fmt_data(data_envio),
        "tem_recurso":      tem_recurso,
        "valor_concedido":  valor_concedido,
        "valor_declarado":  valor_declarado,
        "bolsistas":        bolsistas,
        "atividades":       atividades,
        "total_horas":      total_horas,
        "soma_horas_real":  soma_horas,
        "pub_int":          pub_int,
        "pub_ext":          pub_ext,
        "secoes":           secoes,
        "secao_removida":   secao_removida,
        "produtos":         produtos,
        "prestacao_contas": prestacao_contas,
    }


# ══════════════════════════════════════════════════════════════════════════
# GERAÇÃO DO EDITAL OFICIAL (PDF e EXCEL)
# ══════════════════════════════════════════════════════════════════════════

def gerar_edital_excel(lista_bolsistas_edital, lista_espera_edital, out_dir):
    wb = openpyxl.Workbook()
    
    # --- PLANILHA 1: BOLSISTAS SELECIONADOS ---
    ws = wb.active
    ws.title = "Bolsistas Selecionados"
    
    headers = ["Nome do(a) acadêmico(a)", "RGA", "Projeto vinculado / Coordenação", "Vigência da Bolsa", "Valor da Bolsa", "Fonte pagadora"]
    ws.append(headers)
    
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="1F497D")
    center_align = Alignment(horizontal="center", vertical="center")
    wrap_align = Alignment(vertical="center", wrap_text=True)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    
    for col_idx, cell in enumerate(ws[1], 1):
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border
        
    for b in lista_bolsistas_edital:
        row = [
            b['nome'],
            b['rga'],
            f"{b['projeto']} / {b['coordenador']}",
            f"{b['inicio']} a {b['fim']}",
            f"R$ {b['valor']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            b['fonte']
        ]
        ws.append(row)
        
    for row in ws.iter_rows(min_row=2, max_col=6):
        for idx, cell in enumerate(row):
            cell.border = thin_border
            if idx in (0, 2):
                cell.alignment = wrap_align
            else:
                cell.alignment = center_align
                
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 40
    ws.column_dimensions['D'].width = 22
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15

    # --- PLANILHA 2: LISTA DE ESPERA ---
    ws2 = wb.create_sheet(title="Lista de Espera")
    headers2 = ["Nome do(a) acadêmico(a)", "RGA", "Projeto vinculado / Coordenação", "Classificação"]
    ws2.append(headers2)
    
    for col_idx, cell in enumerate(ws2[1], 1):
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align
        cell.border = thin_border
        
    for b in lista_espera_edital:
        row = [
            b['nome'],
            b['rga'],
            f"{b['projeto']} / {b['coordenador']}",
            b['classificacao']
        ]
        ws2.append(row)
        
    for row in ws2.iter_rows(min_row=2, max_col=4):
        for idx, cell in enumerate(row):
            cell.border = thin_border
            if idx in (0, 2):
                cell.alignment = wrap_align
            else:
                cell.alignment = center_align
                
    ws2.column_dimensions['A'].width = 25
    ws2.column_dimensions['B'].width = 15
    ws2.column_dimensions['C'].width = 40
    ws2.column_dimensions['D'].width = 15

    wb.save(out_dir / "EDITAL_190_2026_PROECE_BOLSISTAS.xlsx")

def gerar_edital_pdf(lista_bolsistas_edital, lista_espera_edital, out_dir):
    story = []
    W = A4[0] - 2 * cm   # margens menores 1cm
    
    E = _estilos_pdf()
    
    story.append(Paragraph("Serviço Público Federal<br/>Ministério da Educação<br/>Fundação Universidade Federal de Mato Grosso do Sul", E["inst"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("EDITAL Nº 190/2026-PROECE/UFMS", E["title"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Resultado do Cadastro de Candidatos a Bolsas de Extensão Universitária - 2026", E["inst"]))
    story.append(Spacer(1, 10))
    
    intro_text = (
        "A FUNDAÇÃO UNIVERSIDADE FEDERAL DE MATO GROSSO DO SUL, por intermédio da Pró-Reitoria de "
        "Extensão, Cultura e Esporte - Proece, no uso de suas atribuições legais e em conformidade com as disposições das "
        "Normas Regulamentadoras das Ações de Extensão da UFMS, aprovada pela Resolução Coex nº 470, de 18 de junho de "
        "2025, as Diretrizes para concessão, acompanhamento e controle de Bolsas e Auxílios na UFMS, aprovada pela Resolução "
        "CD nº 648, de 27 de novembro de 2025, o Decreto nº 7.416, de 30 de dezembro de 2010, a Resolução nº 7, de 18 de "
        "dezembro de 2018, da Câmara de Educação Superior e demais legislações pertinentes, torna público o Resultado do "
        "Cadastro de candidatos a Bolsas de Extensão Universitária - 2026 e a Retificação do Edital nº. 184/2026 - Proece/UFMS."
    )
    story.append(Paragraph(intro_text, E["body"]))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("1. DOS BOLSISTAS SELECIONADOS", E["h1"]))
    
    table_data = [["Nome do(a)\\nacadêmico(a)", "RGA", "Projeto vinculado /\\nCoordenação", "Vigência da\\nBolsa", "Valor da\\nBolsa", "Fonte\\npagadora"]]
    
    for b in lista_bolsistas_edital:
        table_data.append([
            Paragraph(b['nome'], E['small']),
            Paragraph(b['rga'], E['small']),
            Paragraph(f"{b['projeto']} /<br/>{b['coordenador']}", E['small']),
            Paragraph(f"{b['inicio']} a<br/>{b['fim']}", E['small']),
            Paragraph(f"R$ {b['valor']:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."), E['small']),
            Paragraph(b['fonte'], E['small'])
        ])
        
    t = Table(table_data, colWidths=[3.5*cm, 2.5*cm, 6*cm, 3*cm, 2*cm, 2*cm])
    style = [
        ("FONTSIZE",     (0, 0), (-1, -1), 8),
        ("GRID",         (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",        (0, 0), (-1, 0), "CENTER"),
        ("BACKGROUND",   (0, 0), (-1, 0), colors.HexColor("#D9E1F2")),
        ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
    ]
    t.setStyle(TableStyle(style))
    story.append(t)
    
    story.append(Spacer(1, 15))
    story.append(Paragraph("2. DA LISTA DE ESPERA", E["h1"]))
    
    wait_data = [["Nome do(a)\\nacadêmico(a)", "RGA", "Projeto vinculado / Coordenação", "Classificação"]]
    for e in lista_espera_edital:
        wait_data.append([
            Paragraph(e['nome'], E['small']), 
            Paragraph(e['rga'], E['small']), 
            Paragraph(f"{e['projeto']} / {e['coordenador']}", E['small']), 
            Paragraph(e['classificacao'], E['small'])
        ])

    t2 = Table(wait_data, colWidths=[4.5*cm, 2.5*cm, 10*cm, 2*cm])
    t2.setStyle(TableStyle(style))
    story.append(t2)
    
    story.append(Spacer(1, 15))
    story.append(Paragraph("3. DA RETIFICAÇÃO DO EDITAL Nº 184/2026-PROECE/UFMS", E["h1"]))
    
    retif_data = [["Onde se lê:", "Leia-se:"], ["Item 4.2: até o dia 10", "Item 4.2: até o dia 15"]]
    t3 = Table(retif_data, colWidths=[9.5*cm, 9.5*cm])
    t3.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#D9E1F2")),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
    ]))
    story.append(t3)
    
    story.append(Spacer(1, 15))
    story.append(Paragraph("4. DAS PROVIDÊNCIAS PARA A CONCESSÃO DAS BOLSAS", E["h1"]))
    prov = (
        "4.1. O acadêmico selecionado deverá entrar em contato com a Coordenação do Projeto para o envio da documentação.<br/>"
        "4.1.1. As bolsas pagas pela UFMS o acadêmico selecionado deverá enviar à Coordenação o Termo de Concessão (anexo) preenchido e assinado.<br/>"
        "4.1.2. As bolsas pagas pela Fapec, a Coordenação deverá preencher o Termo no processo SEI vinculado ao Projeto.<br/>"
        "4.2. O acadêmico selecionado que desistir da Bolsa deverá formular por escrito o pedido de cancelamento da bolsa, enviar por e-mail para sepex.proece@ufms.br e à Coordenação da Ação de Extensão, antes do dia 15 de cada mês."
    )
    story.append(Paragraph(prov, E["body"]))

    story.append(Spacer(1, 30))
    story.append(Paragraph("Campo Grande, 01 de abril de 2026.", E["body"]))
    story.append(Spacer(1, 20))
    story.append(Paragraph("<b>GABINETE DA PRÓ-REITORIA DE EXTENSÃO, CULTURA E ESPORTE</b><br/>Av Costa e Silva, s/nº - Cidade Universitária<br/>Fone: (67)3345-7232/7233<br/>CEP 79070-900 - Campo Grande - MS", E["inst"]))
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
    story.append(Paragraph("Referência: Processo nº 23104.037776/2025-89 | SEI nº 6329427", E["inst"]))
    doc = SimpleDocTemplate(
        str(out_dir / "EDITAL_190_2026_PROECE_BOLSISTAS.pdf"), pagesize=A4,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
        leftMargin=1 * cm, rightMargin=1 * cm,
    )
    doc.build(story)


# ══════════════════════════════════════════════════════════════════════════
# EXPORTAÇÃO CSV E PREPARAÇÃO DOS DADOS
# ══════════════════════════════════════════════════════════════════════════

def gerar_csvs(lista_dados, lista_arquivos, out_dir):
    # Gerador pseudo-aleatório local para preencher atributos secundários
    rng_csv = random.Random(42)

    lista_bolsistas_edital = []

    # 1. Geração do Gabarito
    with open(out_dir / "gabarito.csv", "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["arquivo", "id_relatorio", "protocolo", "tipo_relatorio", "tipo_erro", "categoria_esperada", "descricao_erro"])
        
        for d, arq in zip(lista_dados, lista_arquivos):
            descricao = ERROS[d["tipo_erro"]][1]
            if d["tipo_erro"] == "E1":
                descricao = f"Valor declarado (R$ {d['valor_declarado']:,.2f}) diverge do concedido (R$ {d['valor_concedido']:,.2f}) em mais de 1%"
            
            writer.writerow([arq, d["id_relatorio"], d["protocolo"], d["tipo_relatorio"], d["tipo_erro"], d["decisao_esperada"], descricao])

    # 2. Geração da Base de Fomentos
    with open(out_dir / "fomentos_concedidos.csv", "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["protocolo_projeto", "valor_aprovado", "fonte", "modalidade_fomento", "data_aprovacao"])
        
        for d in lista_dados:
            if d["tem_recurso"]:
                fonte = rng_csv.choice(["FUNDECT", "CAPES", "CNPq", "PROECE"])
                modalidade = rng_csv.choice(["capital", "custeio", "bolsa"])
                writer.writerow([d["protocolo"], f"{d['valor_concedido']:.2f}", fonte, modalidade, d["inicio_projeto"]])

    # 3. Preparação da Base de Bolsistas Selecionados (Omitindo o arquivo CSV original e usando Edital)
    for d in lista_dados:
        for b in d["bolsistas"]:
            if b["obs"] == "[CPF NAO LOCALIZADO NA BASE OFICIAL]":
                continue # Omitido do edital oficial (O Coordenador falsificou no relatório)
                
            inicio_oficial = b["inicio"]
            fim_oficial = b["fim"]
            
            if b["obs"] == "[PERIODO FORA DO VINCULO FORMAL]":
                ini_d = date(int(b["inicio"][-4:]), int(b["inicio"][3:5]), int(b["inicio"][:2]))
                fim_d = date(int(b["fim"][-4:]), int(b["fim"][3:5]), int(b["fim"][:2]))
                
                diff_months = rng_csv.randint(1, 12)
                sign = rng_csv.choice([-1, 1])
                ini_d = ini_d + timedelta(days=sign * diff_months * 30)
                fim_d = fim_d + timedelta(days=sign * diff_months * 30)
                
                inicio_oficial = ini_d.strftime("%d/%m/%Y")
                fim_oficial = fim_d.strftime("%d/%m/%Y")
            
            lista_bolsistas_edital.append({
                "nome": b["nome"],
                "rga": rga_fake(rng_csv),
                "projeto": d["titulo"],
                "coordenador": d["coordenador"],
                "inicio": inicio_oficial,
                "fim": fim_oficial,
                "valor": rng_csv.choice([700.0, 1000.0, 1200.0]),
                "fonte": rng_csv.choice(["Fapec", "UFMS", "Fundação RTVE"])
            })

    # 4. Geração do CSV do Edital de Bolsistas
    with open(out_dir / "edital_bolsistas.csv", "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Nome do(a) acadêmico(a)", "RGA", "Projeto vinculado / Coordenação", "Vigência da Bolsa", "Valor da Bolsa", "Fonte pagadora"])
        
        for b in lista_bolsistas_edital:
            vigencia = f"{b['inicio']} a {b['fim']}"
            writer.writerow([b['nome'], b['rga'], f"{b['projeto']} / {b['coordenador']}", vigencia, b['valor'], b['fonte']])

    # 5. Geração de Bolsistas de Espera
    lista_espera_edital = []
    NOMES_ESPERA = ["Mateus Silva", "Larissa Souza", "Pedro Henrique Alves", "Camila Rodrigues", "João Pedro Costa"]
    for idx_espera, n_espera in enumerate(NOMES_ESPERA, start=1):
        lista_espera_edital.append({
            "nome": n_espera,
            "rga": rga_fake(rng_csv),
            "projeto": rng_csv.choice(TITULOS),
            "coordenador": rng_csv.choice(NOMES),
            "classificacao": f"{idx_espera}º"
        })
                
    # 6. Geração do Edital em PDF e Excel
    gerar_edital_pdf(lista_bolsistas_edital, lista_espera_edital, out_dir)
    gerar_edital_excel(lista_bolsistas_edital, lista_espera_edital, out_dir)


# ══════════════════════════════════════════════════════════════════════════
# GERADOR DOCX — Formato SIGProj/UFMS (tabelas, cabeçalho institucional)
# ══════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    tc  = cell._tc
    tcp = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcp.append(shd)

def _style_cell(cell, bold=False, bg=None, size=10):
    if bg:
        _set_cell_bg(cell, bg)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.bold = bold
            run.font.size = Pt(size)

def _add_row(table, label, value, bg_label="D9E1F2"):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)
    _style_cell(row.cells[0], bold=True, bg=bg_label)
    _style_cell(row.cells[1])

def _make_table(doc, col_widths=(Cm(6), Cm(11))):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.columns[0].width = col_widths[0]
    t.columns[1].width = col_widths[1]
    return t

def gerar_docx(dados: dict, path: Path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2)

    # Cabeçalho
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = hdr.add_run("Servico Publico Federal – Ministerio da Educacao\n")
    r1.bold = True; r1.font.size = Pt(11)
    r2 = hdr.add_run("Fundacao Universidade Federal de Mato Grosso do Sul")
    r2.bold = True; r2.font.size = Pt(11)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(f"RELATORIO {dados['tipo_relatorio'].upper()} DE ACAO DE EXTENSAO")
    tr.bold = True; tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    # Dados da submissão
    doc.add_heading("Dados da Submissao", level=2)
    t = _make_table(doc)
    for label, valor in [
        ("Titulo da Submissao",           dados["titulo"]),
        ("Coordenador",                   f"{dados['coordenador']} — CPF: {dados['cpf_coord']}"),
        ("Protocolo",                     dados["protocolo"]),
        ("Tipo de Relatorio",             dados["tipo_relatorio"]),
        ("Numero do Relatorio",           dados["id_relatorio"]),
        ("Data de inicio",                dados["inicio_projeto"]),
        ("Data de encerramento prevista", dados["fim_projeto"]),
        ("Data de envio",                 dados["data_envio"]),
        ("Unidade executora",             dados["unidade"]),
        ("Area do Conhecimento",          dados["area"]),
        ("Linha de Extensao",             dados["linha"]),
        ("Situacao",                      "Enviado"),
    ]:
        _add_row(t, label, valor)

    doc.add_paragraph()

    # ODS
    doc.add_heading("Objetivos de Desenvolvimento Sustentavel (ODS)", level=2)
    for o in dados["ods"]:
        doc.add_paragraph(o, style="List Bullet").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Objetivo
    doc.add_heading("Objetivo do Projeto", level=2)
    doc.add_paragraph(dados["secoes"]["objetivos_geral"])
    doc.add_paragraph()

    # Público interno
    doc.add_heading("Publico Envolvido", level=2)
    pi = doc.add_paragraph("Interno (UFMS):")
    pi.runs[0].bold = True
    t2 = _make_table(doc)
    for cat, qtd in dados["pub_int"].items():
        _add_row(t2, cat, str(qtd))

    pe = doc.add_paragraph("Externo:")
    pe.runs[0].bold = True
    t3 = _make_table(doc)
    for cat, qtd in dados["pub_ext"].items():
        _add_row(t3, cat, str(qtd))
    doc.add_paragraph()

    # Recursos financeiros
    doc.add_heading("Recursos Financeiros", level=2)
    t4 = _make_table(doc)
    _add_row(t4, "Projeto teve recursos financeiros?",
             "Sim" if dados["tem_recurso"] else "Nao")
    if dados["tem_recurso"]:
        _add_row(t4, "Valor total declarado (R$)", f"{dados['valor_declarado']:,.2f}")
    doc.add_paragraph()

    if dados["tem_recurso"] and dados["prestacao_contas"]:
        doc.add_heading("Prestacao de Contas", level=2)
        tpc = doc.add_table(rows=1, cols=2)
        tpc.style = "Table Grid"
        tpc.rows[0].cells[0].text = "Item de Despesa"
        tpc.rows[0].cells[1].text = "Valor (R$)"
        _style_cell(tpc.rows[0].cells[0], bold=True, bg="D9E1F2")
        _style_cell(tpc.rows[0].cells[1], bold=True, bg="D9E1F2")
        for pc in dados["prestacao_contas"]:
            row = tpc.add_row()
            row.cells[0].text = pc["item"]
            row.cells[1].text = f"{pc['valor']:,.2f}"
            _style_cell(row.cells[0])
            _style_cell(row.cells[1])
        doc.add_paragraph()

    # Bolsistas
    doc.add_heading("Bolsistas / Voluntarios", level=2)
    for b in dados["bolsistas"]:
        tb = _make_table(doc)
        _add_row(tb, "Nome",             b["nome"])
        _add_row(tb, "CPF",              b["cpf"])
        _add_row(tb, "Tipo de vinculo",  b["vinculo"])
        _add_row(tb, "Periodo declarado", f"{b['inicio']} a {b['fim']}")
        if b["obs"]:
            _add_row(tb, "Observacao", b["obs"], bg_label="FFD7D7")
        doc.add_paragraph()

    # Atividades
    doc.add_heading("Atividades Realizadas", level=2)
    ta = doc.add_table(rows=1, cols=4)
    ta.style = "Table Grid"
    for i, h in enumerate(["Descricao", "Data", "Local", "CH (h)"]):
        c = ta.rows[0].cells[i]
        c.text = h
        _style_cell(c, bold=True, bg="D9E1F2")
    for a in dados["atividades"]:
        row = ta.add_row()
        row.cells[0].text = a["descricao"]
        row.cells[1].text = a["data"]
        row.cells[2].text = a["local"]
        row.cells[3].text = str(a["horas"])
        for c in row.cells:
            _style_cell(c)

    tot_p = doc.add_paragraph()
    rt = tot_p.add_run(f"Total de horas declarado: {dados['total_horas']} h")
    rt.bold = True; rt.font.size = Pt(10)
    if dados["tipo_erro"] == "E5":
        re = tot_p.add_run(
            f"  [INCONSISTENCIA: soma das atividades = {dados['soma_horas_real']} h]"
        )
        re.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        re.font.size = Pt(9)
    doc.add_paragraph()

    # Seções textuais
    for chave, titulo_sec in [
        ("sintese",       "Sintese da Execucao e dos Resultados Obtidos"),
        ("resultados",    "Resultados Alcancados"),
        ("metodologia",   "Metodologia Utilizada"),
        ("dificuldades",  "Dificuldades Encontradas"),
        ("consideracoes", "Consideracoes Finais"),
    ]:
        doc.add_heading(titulo_sec, level=2)
        conteudo = dados["secoes"].get(chave, "")
        if not conteudo:
            p = doc.add_paragraph("[SECAO AUSENTE]")
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(conteudo)
        doc.add_paragraph()

    # Produtos
    doc.add_heading("Resultados Alcancados (Produtos)", level=2)
    tp2 = doc.add_table(rows=1, cols=3)
    tp2.style = "Table Grid"
    for i, h in enumerate(["Tipo de Entrega", "Titulo", "Ano"]):
        c = tp2.rows[0].cells[i]
        c.text = h
        _style_cell(c, bold=True, bg="D9E1F2")
    for prod in dados["produtos"]:
        row = tp2.add_row()
        row.cells[0].text = prod["tipo"]
        row.cells[1].text = prod["titulo"]
        row.cells[2].text = prod["ano"]
        for c in row.cells:
            _style_cell(c)
    doc.add_paragraph()

    # Referências
    doc.add_heading("Referencias Bibliograficas", level=2)
    for ref in dados["secoes"]["referencias"]:
        doc.add_paragraph(ref, style="List Bullet").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Assinatura
    doc.add_paragraph(
        f"Campo Grande – MS, {dados['data_envio']}\n\n"
        "_______________________________________\n"
        f"{dados['coordenador']} – Coordenador(a)\n\n"
        "_______________________________________\n"
        "Coordenacao de Extensao"
    ).runs[0].font.size = Pt(10)

    doc.save(path)


# ══════════════════════════════════════════════════════════════════════════
# GERADOR PDF — Formato IFMG (discursivo, seções numeradas)
# ══════════════════════════════════════════════════════════════════════════

def _estilos_pdf():
    base = getSampleStyleSheet()
    def S(name, **kw):
        return ParagraphStyle(name, parent=base["Normal"], **kw)

    return {
        "inst":   S("inst",  fontSize=10, leading=13, alignment=TA_CENTER,
                    fontName="Helvetica-Bold", spaceAfter=2),
        "title":  S("title", fontSize=13, leading=16, alignment=TA_CENTER,
                    fontName="Helvetica-Bold", spaceAfter=4),
        "h1":     S("h1",    fontSize=11, leading=14, fontName="Helvetica-Bold",
                    spaceBefore=14, spaceAfter=4,
                    textColor=colors.HexColor("#1F497D")),
        "h2":     S("h2",    fontSize=10, leading=13, fontName="Helvetica-Bold",
                    spaceBefore=8, spaceAfter=3),
        "body":   S("body",  fontSize=10, leading=14, alignment=TA_JUSTIFY,
                    spaceAfter=6),
        "bullet": S("bullet",fontSize=10, leading=14, leftIndent=14,
                    spaceAfter=3),
        "error":  S("error", fontSize=10, leading=13, fontName="Helvetica-Bold",
                    textColor=colors.red),
        "small":  S("small", fontSize=9,  leading=12, spaceAfter=2),
    }


def _tab(rows, col_widths, header=True):
    t = Table(rows, colWidths=col_widths)
    style = [
        ("FONTSIZE",     (0, 0), (-1, -1), 9),
        ("GRID",         (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING",   (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
    ]
    if header:
        style += [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9E1F2")),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ]
    t.setStyle(TableStyle(style))
    return t


def gerar_pdf(dados: dict, path: Path):
    E = _estilos_pdf()
    story = []
    W = A4[0] - 5 * cm   # largura útil (margens 3cm + 2cm)

    # Cabeçalho
    story += [
        Paragraph("Servico Publico Federal – Ministerio da Educacao", E["inst"]),
        Paragraph("Fundacao Universidade Federal de Mato Grosso do Sul", E["inst"]),
        Paragraph(f"{dados['unidade']} – Coordenacao de Extensao", E["inst"]),
        Spacer(1, 6),
        HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1F497D")),
        Spacer(1, 6),
        Paragraph(dados["titulo"], E["title"]),
        Paragraph(
            f"Relatorio {dados['tipo_relatorio']} — Protocolo: {dados['protocolo']}",
            E["inst"],
        ),
        HRFlowable(width="100%", thickness=0.5, color=colors.grey),
        Spacer(1, 10),
    ]

    # Ficha de identificação
    story.append(Paragraph("Identificacao", E["h1"]))
    ficha = [
        ["Campo", "Informacao"],
        ["Coordenador", dados["coordenador"]],
        ["CPF", dados["cpf_coord"]],
        ["No do Relatorio", dados["id_relatorio"]],
        ["Tipo de Relatorio", dados["tipo_relatorio"]],
        ["Area do Conhecimento", dados["area"]],
        ["Linha de Extensao", dados["linha"]],
        ["Inicio do projeto", dados["inicio_projeto"]],
        ["Encerramento previsto", dados["fim_projeto"]],
        ["Data de envio", dados["data_envio"]],
        ["Situacao", "Enviado"],
    ]
    story.append(_tab(ficha, [5 * cm, W - 5 * cm]))
    story.append(Spacer(1, 10))

    # ODS
    story.append(Paragraph("Objetivos de Desenvolvimento Sustentavel", E["h1"]))
    for o in dados["ods"]:
        story.append(Paragraph(f"• {o}", E["bullet"]))
    story.append(Spacer(1, 8))

    # Objetivo
    story.append(Paragraph("Objetivo do Projeto", E["h1"]))
    story.append(Paragraph(dados["secoes"]["objetivos_geral"], E["body"]))
    story.append(Spacer(1, 10))

    # Público
    story.append(Paragraph("Publico Envolvido", E["h1"]))
    story.append(Paragraph("Interno (UFMS)", E["h2"]))
    story.append(_tab(
        [["Categoria", "Qtd."]] + [[k, str(v)] for k, v in dados["pub_int"].items()],
        [10 * cm, W - 10 * cm],
    ))
    story.append(Spacer(1, 4))
    story.append(Paragraph("Externo", E["h2"]))
    story.append(_tab(
        [["Categoria", "Qtd."]] + [[k, str(v)] for k, v in dados["pub_ext"].items()],
        [10 * cm, W - 10 * cm],
    ))
    story.append(Spacer(1, 10))

    # Recursos
    story.append(Paragraph("Recursos Financeiros", E["h1"]))
    rec = [["Projeto teve recursos financeiros?",
            "Sim" if dados["tem_recurso"] else "Nao"]]
    if dados["tem_recurso"]:
        rec.append(["Valor total declarado (R$)", f"{dados['valor_declarado']:,.2f}"])
    story.append(_tab(rec, [9 * cm, W - 9 * cm], header=False))
    story.append(Spacer(1, 10))

    if dados["tem_recurso"] and dados["prestacao_contas"]:
        story.append(Paragraph("Prestacao de Contas", E["h1"]))
        pc_rows = [["Item de Despesa", "Valor (R$)"]]
        for pc in dados["prestacao_contas"]:
            pc_rows.append([pc["item"], f"{pc['valor']:,.2f}"])
        story.append(_tab(pc_rows, [10 * cm, W - 10 * cm]))
        story.append(Spacer(1, 10))

    # Bolsistas
    story.append(Paragraph("Bolsistas / Voluntarios", E["h1"]))
    bol_rows = [["Nome", "CPF", "Vinculo", "Periodo declarado"]]
    for b in dados["bolsistas"]:
        bol_rows.append([b["nome"], b["cpf"], b["vinculo"],
                          f"{b['inicio']} a {b['fim']}"])
    story.append(_tab(bol_rows, [4 * cm, 3.2 * cm, 2.3 * cm, W - 9.5 * cm]))
    for b in dados["bolsistas"]:
        if b["obs"]:
            story.append(Paragraph(f"Atencao: {b['nome']} — {b['obs']}", E["error"]))
    story.append(Spacer(1, 10))

    # Atividades
    story.append(Paragraph("Atividades Realizadas", E["h1"]))
    atv_rows = [["Descricao", "Data", "Local", "CH (h)"]]
    for a in dados["atividades"]:
        atv_rows.append([a["descricao"], a["data"], a["local"], str(a["horas"])])
    story.append(_tab(atv_rows, [6 * cm, 2.5 * cm, 4 * cm, W - 12.5 * cm]))
    txt_h = f"Total de horas declarado: {dados['total_horas']} h"
    if dados["tipo_erro"] == "E5":
        txt_h += f"  [INCONSISTENCIA: soma = {dados['soma_horas_real']} h]"
        story.append(Paragraph(txt_h, E["error"]))
    else:
        story.append(Paragraph(f"<b>{txt_h}</b>", E["body"]))
    story.append(Spacer(1, 10))

    # Seções textuais numeradas
    for num, chave, titulo_sec in [
        ("1", "sintese",       "Sintese da Execucao e dos Resultados Obtidos"),
        ("2", "resultados",    "Resultados Alcancados"),
        ("3", "metodologia",   "Metodologia Utilizada"),
        ("4", "contribuicoes", "Contribuicoes para o Ensino, Pesquisa e Extensao"),
        ("5", "dificuldades",  "Dificuldades Encontradas"),
        ("6", "consideracoes", "Consideracoes Finais"),
    ]:
        story.append(Paragraph(f"{num}  {titulo_sec}", E["h1"]))
        conteudo = dados["secoes"].get(chave, "")
        if not conteudo:
            story.append(Paragraph("[SECAO AUSENTE]", E["error"]))
        else:
            story.append(Paragraph(conteudo, E["body"]))
        story.append(Spacer(1, 6))

    # Produtos
    story.append(Paragraph("Produtos / Entregas", E["h1"]))
    prod_rows = [["Tipo de Entrega", "Titulo", "Ano"]] + [
        [p["tipo"], p["titulo"], p["ano"]] for p in dados["produtos"]
    ]
    story.append(_tab(prod_rows, [4.5 * cm, W - 6.5 * cm, 2 * cm]))
    story.append(Spacer(1, 10))

    # Referências
    story.append(Paragraph("Referencias Bibliograficas", E["h1"]))
    for ref in dados["secoes"]["referencias"]:
        story.append(Paragraph(f"• {ref}", E["bullet"]))
    story.append(Spacer(1, 14))

    # Assinatura
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Spacer(1, 6))
    ass = [
        [f"Campo Grande – MS, {dados['data_envio']}", ""],
        ["", ""],
        [f"{dados['coordenador']}", "Coordenacao de Extensao"],
        ["Coordenador(a)", ""],
    ]
    story.append(_tab(ass, [W / 2, W / 2], header=False))

    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=3 * cm, rightMargin=2 * cm,
    )
    doc.build(story)


# ══════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Gera 30 relatorios sinteticos (15 DOCX + 15 PDF) para o Processo 3.4"
    )
    parser.add_argument("--seed", type=int, default=42)
    parser.add_argument("--out",  type=str, default="./relatorios")
    args = parser.parse_args()

    rng = random.Random(args.seed)
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)

    # 30 tipos de erro conforme distribuição definida
    tipos = []
    for codigo, (qtd, _) in ERROS.items():
        tipos.extend([codigo] * qtd)
    rng.shuffle(tipos)

    # 15 DOCX + 15 PDF, embaralhados
    formatos = ["docx"] * 15 + ["pdf"] * 15
    rng.shuffle(formatos)

    print(f"Gerando 30 relatorios em: {out.resolve()}\n")
    print(f"{'#':>3}  {'Fmt':<5}  {'Erro':<4}  {'Decisao esperada':<32}  Arquivo")
    print("─" * 82)

    lista_dados_gerados = []
    lista_arquivos_gerados = []

    for idx, (tipo_erro, fmt) in enumerate(zip(tipos, formatos), start=1):
        dados   = gerar_dados(rng, idx, tipo_erro)
        arquivo = f"relatorio_{idx:02d}_{tipo_erro}.{fmt}"
        caminho = out / arquivo

        lista_dados_gerados.append(dados)
        lista_arquivos_gerados.append(arquivo)

        if fmt == "docx":
            gerar_docx(dados, caminho)
        else:
            gerar_pdf(dados, caminho)

        print(f"{idx:>3}  {fmt:<5}  {tipo_erro:<4}  {dados['decisao_esperada']:<32}  {arquivo}")

    print("\n" + "═" * 82)
    print(f"30 relatorios gerados em: {out.resolve()}")
    print("═" * 82)
    print("\nDistribuicao de erros:")
    for cod, (qtd, desc) in ERROS.items():
        print(f"  {cod}: {qtd:>2}x  —  {desc}")
    print("\nFormatos: 15 DOCX (estilo SIGProj/UFMS)  |  15 PDF (estilo IFMG discursivo)")

    # Geração dos arquivos CSV
    gerar_csvs(lista_dados_gerados, lista_arquivos_gerados, out)
    print("\n[+] Bases de dados exportadas com sucesso na mesma pasta:")
    print("  - fomentos_concedidos.csv")
    print("  - edital_bolsistas.csv")
    print("  - gabarito.csv")
    print("\n[+] Edital Oficial gerado:")
    print("  - EDITAL_190_2026_PROECE_BOLSISTAS.pdf")
    print("  - EDITAL_190_2026_PROECE_BOLSISTAS.xlsx\n")


if __name__ == "__main__":
    main()